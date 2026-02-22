"""Expense Word Generation Module - Receipt proof documents."""
import os
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Word output directory
OUTPUT_DIR = os.path.expanduser("~/projects/bim/expenses/output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Image width from Dalux template
IMAGE_WIDTH_INCHES = 4.6


def generate_word(month_key: str, receipts: List[Dict[str, Any]]) -> str:
    """
    Generate Word document with receipt images.
    
    Args:
        month_key: YYYY-MM format
        receipts: List of receipt dictionaries
        
    Returns:
        Path to generated Word file
    """
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Expense Receipts - {month_key}")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Sort by date
    receipts_sorted = sorted(
        [r for r in receipts if r.get('status') == 'active'],
        key=lambda r: r.get('date', '')
    )
    
    for receipt in receipts_sorted:
        # Label with description
        vendor = receipt.get('vendor', 'Unknown')
        desc = receipt.get('description', '')
        date = receipt.get('date', '')
        amount = receipt.get('amount_eur', 0)
        currency = receipt.get('original_currency', '')
        
        label_text = f"{date} - {vendor}"
        if desc:
            label_text += f": {desc}"
        label_text += f" ({amount:.2f} EUR)"
        
        label = doc.add_paragraph(label_text)
        label.runs[0].font.bold = True
        
        # Add receipt image if exists
        image_path = receipt.get('image_path', '')
        if image_path and os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(IMAGE_WIDTH_INCHES))
            except Exception as e:
                doc.add_paragraph(f"[Image not available: {image_path}]")
        else:
            doc.add_paragraph("[No receipt image]")
        
        # Spacing between receipts
        doc.add_paragraph()
    
    # Save
    year, month = month_key.split('-')
    month_name = get_month_name(month)
    filename = f"RZE_-_Travel_and_Expenses_Dalux_{month_name}_{year}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    return filepath


def get_month_name(month_num: str) -> str:
    """Convert month number to name."""
    months = ['JAN', 'FEB', 'MAR', 'APR', 'MAY', 'JUN', 
              'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC']
    try:
        return months[int(month_num) - 1]
    except (ValueError, IndexError):
        return month_num


if __name__ == "__main__":
    # Test with sample data
    test_receipts = [
        {
            'date': '2026-02-01',
            'vendor': 'Uber',
            'description': 'Airport ride',
            'amount_eur': 23.00,
            'original_currency': 'USD',
            'image_path': '',
            'status': 'active'
        }
    ]
    
    filepath = generate_word('2026-02', test_receipts)
    print(f"Generated: {filepath}")
